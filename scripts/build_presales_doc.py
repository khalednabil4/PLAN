from __future__ import annotations

import argparse
import json
import math
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
SOURCE_PATH = ROOT / "docs" / "presales_checklist_source.json"
OUTPUT_PATH = ROOT / "docs" / "Presales_Department_Flow.docx"
GENERATED_LOGO_PATH = ROOT / "docs" / "assets" / "alnafitha_brand_mark_generated.png"


def hex_to_rgb(value: str) -> tuple[int, int, int]:
    cleaned = value.lstrip("#")
    return tuple(int(cleaned[index : index + 2], 16) for index in (0, 2, 4))


def apply_cell_fill(cell, hex_color: str) -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color.lstrip("#"))
    cell_properties.append(shading)


def set_cell_margins(cell, top: int, start: int, bottom: int, end: int) -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    margins = cell_properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        cell_properties.append(margins)

    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = margins.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, size: int, color: str | None = None, bold: bool = False) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*hex_to_rgb(color))


def add_bullets(document: Document, items: list[str], color: str) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        set_run_font(run, 10, color=color)
        paragraph.paragraph_format.space_after = Pt(4)


def add_body_paragraph(document: Document, text: str, color: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, 10, color=color)


def create_generated_logo(path: Path, brand: dict[str, str], company_name: str) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)

    width, height = 1400, 520
    primary = hex_to_rgb(brand["primary"])
    secondary = hex_to_rgb(brand["secondary"])
    accent = hex_to_rgb(brand["accent"])

    image = Image.new("RGB", (width, height), primary)
    draw = ImageDraw.Draw(image)

    burst_center = (1100, 260)
    burst_outer = 160
    burst_inner = 58
    for degree in range(0, 360, 10):
        radians = degree * math.pi / 180
        x1 = burst_center[0] + int(burst_inner * math.cos(radians))
        y1 = burst_center[1] + int(burst_inner * math.sin(radians))
        x2 = burst_center[0] + int(burst_outer * math.cos(radians))
        y2 = burst_center[1] + int(burst_outer * math.sin(radians))
        draw.line((x1, y1, x2, y2), fill=accent, width=4)

    draw.ellipse(
        (
            burst_center[0] - 88,
            burst_center[1] - 88,
            burst_center[0] + 88,
            burst_center[1] + 88,
        ),
        fill=secondary,
    )

    font_paths = [
        Path("C:/Windows/Fonts/calibri.ttf"),
        Path("C:/Windows/Fonts/arial.ttf"),
    ]
    selected_font = next((font for font in font_paths if font.exists()), None)
    if selected_font:
        name_font = ImageFont.truetype(str(selected_font), 112)
        tag_font = ImageFont.truetype(str(selected_font), 66)
    else:
        name_font = ImageFont.load_default()
        tag_font = ImageFont.load_default()

    draw.text((95, 180), company_name.replace(" IT", ""), fill=(255, 255, 255), font=name_font)

    badge_text = "IT"
    badge_box = draw.textbbox((0, 0), badge_text, font=tag_font)
    badge_width = badge_box[2] - badge_box[0]
    badge_height = badge_box[3] - badge_box[1]
    draw.text(
        (
            burst_center[0] - badge_width / 2,
            burst_center[1] - badge_height / 2 - 6,
        ),
        badge_text,
        fill=(255, 255, 255),
        font=tag_font,
    )

    image.save(path)
    return path


def add_cover(document: Document, data: dict, logo_path: Path) -> None:
    brand = data["brand"]
    metadata = data["document"]

    table = document.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    apply_cell_fill(cell, brand["primary"])
    set_cell_margins(cell, top=160, start=240, bottom=160, end=240)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(logo_path), width=Inches(5.9))
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    document.add_paragraph()

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(metadata["title"])
    set_run_font(title_run, 22, color=brand["primary"], bold=True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(metadata["subtitle"])
    set_run_font(subtitle_run, 11, color=brand["text_dark"])

    meta_table = document.add_table(rows=4, cols=2)
    meta_table.style = "Table Grid"
    metadata_rows = [
        ("Company", metadata["company"]),
        ("Department", metadata["department"]),
        ("Version", metadata["version"]),
        ("Last Updated", metadata["last_updated"]),
    ]

    for index, (label, value) in enumerate(metadata_rows):
        label_cell, value_cell = meta_table.rows[index].cells
        apply_cell_fill(label_cell, brand["muted_fill"])
        set_cell_margins(label_cell, top=100, start=120, bottom=100, end=120)
        set_cell_margins(value_cell, top=100, start=120, bottom=100, end=120)

        label_run = label_cell.paragraphs[0].add_run(label)
        set_run_font(label_run, 10, color=brand["text_dark"], bold=True)

        value_run = value_cell.paragraphs[0].add_run(value)
        set_run_font(value_run, 10, color=brand["text_dark"])

    document.add_page_break()


def add_section_heading(document: Document, text: str, brand: dict[str, str]) -> None:
    heading = document.add_paragraph()
    heading.paragraph_format.space_before = Pt(6)
    heading.paragraph_format.space_after = Pt(6)
    run = heading.add_run(text)
    set_run_font(run, 15, color=brand["primary"], bold=True)


def add_summary_table(document: Document, data: dict) -> None:
    brand = data["brand"]
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    headings = ["Page", "ID", "Checklist Item"]
    for index, label in enumerate(headings):
        cell = table.rows[0].cells[index]
        apply_cell_fill(cell, brand["primary"])
        set_cell_margins(cell, top=100, start=100, bottom=100, end=100)
        run = cell.paragraphs[0].add_run(label)
        set_run_font(run, 10, color=brand["text_light"], bold=True)

    for page in data["pages"]:
        for item in page["checklists"]:
            row = table.add_row().cells
            values = [page["name"], item["id"], item["title"]]
            for index, value in enumerate(values):
                set_cell_margins(row[index], top=100, start=100, bottom=100, end=100)
                run = row[index].paragraphs[0].add_run(value)
                set_run_font(run, 9, color=brand["text_dark"])


def build_document(source_path: Path, output_path: Path, logo_path: Path | None = None) -> Path:
    data = json.loads(source_path.read_text(encoding="utf-8"))
    brand = data["brand"]

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    styles["Normal"].font.size = Pt(10)

    if logo_path:
        resolved_logo_path = logo_path
    else:
        resolved_logo_path = create_generated_logo(GENERATED_LOGO_PATH, brand, data["document"]["company"])

    add_cover(document, data, resolved_logo_path)

    add_section_heading(document, "Checklist Summary", brand)
    add_summary_table(document, data)

    for page in data["pages"]:
        document.add_section(WD_SECTION.NEW_PAGE)
        add_section_heading(document, page["name"], brand)
        add_body_paragraph(document, page["intro"], brand["text_dark"])

        for checklist in page["checklists"]:
            add_section_heading(document, f'{checklist["id"]} - {checklist["title"]}', brand)

            add_section_heading(document, "Checklist Description", brand)
            add_body_paragraph(document, checklist["description"], brand["text_dark"])

            add_section_heading(document, checklist["solution_title"], brand)
            add_bullets(document, checklist["solution_bullets"], brand["text_dark"])

    output_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        document.save(output_path)
        return output_path
    except PermissionError:
        fallback_path = output_path.with_name(f"{output_path.stem}_updated{output_path.suffix}")
        document.save(fallback_path)
        return fallback_path


def extract_document(doc_path: Path) -> str:
    document = Document(doc_path)
    lines: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)

    for table in document.tables:
        lines.append("")
        for row in table.rows:
            cell_text = " | ".join(cell.text.strip().replace("\n", " ") for cell in row.cells)
            if cell_text.strip():
                lines.append(cell_text)

    return "\n".join(lines).strip()


def main() -> None:
    parser = argparse.ArgumentParser(description="Build or read the Presales Word document.")
    subparsers = parser.add_subparsers(dest="command", required=True)

    build_parser = subparsers.add_parser("build", help="Build the Word document from the JSON source.")
    build_parser.add_argument("--source", default=str(SOURCE_PATH))
    build_parser.add_argument("--output", default=str(OUTPUT_PATH))
    build_parser.add_argument("--logo", default=None, help="Optional path to an official company logo image.")

    extract_parser = subparsers.add_parser("extract", help="Extract readable text from an existing Word document.")
    extract_parser.add_argument("doc_path", nargs="?", default=str(OUTPUT_PATH))

    args = parser.parse_args()

    if args.command == "build":
        custom_logo = Path(args.logo) if args.logo else None
        output = build_document(Path(args.source), Path(args.output), custom_logo)
        print(output)
        return

    extracted_text = extract_document(Path(args.doc_path))
    print(extracted_text)


if __name__ == "__main__":
    main()
